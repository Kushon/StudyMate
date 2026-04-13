import io
from docx import Document
from loguru import logger


def parse_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    result = "\n\n".join(parts)
    logger.info(f"DOCX extracted {len(result)} characters, {len(doc.paragraphs)} paragraphs")
    return result
